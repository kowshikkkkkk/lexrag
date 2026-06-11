import hashlib
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from config.constants import (
    SUPPORTED_EXTENSIONS,
    META_SOURCE,
    META_DOC_TYPE,
    META_FILE_HASH,
    META_INGESTED_AT,
    DOC_TYPE_GENERIC,
)
from config.exceptions import UnsupportedFileTypeError, IngestionError
from observability.logger import setup_logger

logger = setup_logger(__name__)


@dataclass
class Document:
    """
    A single loaded document before chunking.
    This object travels from ingestion → chunking → embedding.
    """
    text: str
    metadata: dict = field(default_factory=dict)
    file_path: Optional[str] = None


def compute_file_hash(file_path: Path) -> str:
    """MD5 hash of file contents — used for deduplication."""
    hasher = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def load_txt(file_path: Path) -> str:
    """Load plain text file."""
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return file_path.read_text(encoding="latin-1")


def load_pdf(file_path: Path) -> str:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception as e:
        raise IngestionError(f"Failed to parse PDF {file_path.name}: {e}")


def load_docx(file_path: Path) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        raise IngestionError(f"Failed to parse DOCX {file_path.name}: {e}")


def normalize_text(text: str) -> str:
    """
    Clean raw extracted text.
    - Normalize unicode
    - Collapse excessive whitespace
    - Remove null bytes
    """
    import unicodedata
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\x00", "")
    # Collapse 3+ newlines into 2
    import re
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse multiple spaces into one
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def load_document(file_path: str, doc_type: str = DOC_TYPE_GENERIC) -> Document:
    """
    Main entry point for ingestion.
    Takes a file path, returns a clean Document object.
    """
    path = Path(file_path)

    if not path.exists():
        raise IngestionError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. Supported: {SUPPORTED_EXTENSIONS}"
        )

    logger.info(f"Loading document", extra={"file": path.name, "type": ext})

    # Extract raw text based on file type
    if ext == ".txt":
        raw_text = load_txt(path)
    elif ext == ".pdf":
        raw_text = load_pdf(path)
    elif ext == ".docx":
        raw_text = load_docx(path)

    # Normalize
    clean_text = normalize_text(raw_text)

    if not clean_text:
        raise IngestionError(f"No text could be extracted from {path.name}")

    # Compute hash for deduplication
    file_hash = compute_file_hash(path)

    metadata = {
        META_SOURCE: path.name,
        META_DOC_TYPE: doc_type,
        META_FILE_HASH: file_hash,
        META_INGESTED_AT: datetime.now(timezone.utc).isoformat(),
    }

    logger.info(
        "Document loaded successfully",
        extra={
            "file": path.name,
            "chars": len(clean_text),
            "hash": file_hash,
        },
    )

    return Document(text=clean_text, metadata=metadata, file_path=str(path))